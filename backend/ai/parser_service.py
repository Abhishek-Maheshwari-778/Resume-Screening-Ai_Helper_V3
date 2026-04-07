"""
Robust file text extraction with multi-library fallback chain.
Handles: PDF (standard & scanned), DOCX, TXT, and misnamed files.
Fallback chain: pypdf -> PyPDF2 -> pdfminer -> bruteforce text
"""
import io
import os
import re

#  PDF Libraries (try all) 
PYPDF_AVAILABLE = False
PYPDF2_AVAILABLE = False
PDFMINER_AVAILABLE = False
FITZ_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    pass

try:
    from pypdf import PdfReader as PyPdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    pass

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    pass

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDFMINER_AVAILABLE = True
except ImportError:
    pass

#  DOCX Library 
DOCX_AVAILABLE = False
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    pass


def _clean_text(text: str) -> str:
    """Clean and normalize extracted text for AI processing."""
    if not text:
        return ""
    # Remove control characters and surrogate characters that break JSON
    text = "".join(char for char in text if char.isprintable() or char in "\n\r\t")
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Multi-library fallback chain for PDF extraction.
    pypdf -> fitz -> PyPDF2 -> pdfminer -> brute-force stream recovery
    """
    if not file_bytes:
        return ""

    # 1. Try pypdf (Primary - most stable)
    if PYPDF_AVAILABLE:
        try:
            print(" Attempting pypdf...")
            reader = PyPdfReader(io.BytesIO(file_bytes))
            pages = []
            for page in reader.pages:
                txt = page.extract_text()
                if txt: pages.append(txt)
            text = "\n".join(pages)
            if text and len(text.strip()) > 30:
                print(f" pypdf success: {len(text)} chars")
                return _clean_text(text)
        except Exception as e:
            print(f" pypdf failed: {e}")

    # 2. Try fitz (Secondary - high fidelity)
    if FITZ_AVAILABLE:
        try:
            print(" Attempting fitz deep-scan...")
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                text = "\n".join([page.get_text() for page in doc])
            if text and len(text.strip()) > 30:
                print(f" fitz success: {len(text)} chars")
                return _clean_text(text)
        except Exception as e:
            print(f" fitz failed: {e}")

    # 3. Try PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            print(" Attempting PyPDF2 fallback...")
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = "\n".join([p.extract_text() or "" for p in reader.pages])
            if text and len(text.strip()) > 30:
                print(f" PyPDF2 success: {len(text)} chars")
                return _clean_text(text)
        except Exception as e:
            print(f" PyPDF2 failed: {e}")

    # 4. Try pdfminer
    if PDFMINER_AVAILABLE:
        try:
            print(" Attempting pdfminer fallback...")
            text = pdfminer_extract(io.BytesIO(file_bytes))
            if text and len(text.strip()) > 30:
                print(f" pdfminer success: {len(text)} chars")
                return _clean_text(text)
        except Exception as e:
            print(f" pdfminer failed: {e}")

    # 5. SUPER-FALLBACK: Brute-force stream recovery
    print(" Running PDF brute-force recovery...")
    try:
        raw_text = file_bytes.decode('latin-1', errors='replace')
        # PDF streams often contain text in parentheses
        parts = re.findall(r'\((.*?)\)', raw_text)
        candidate = " ".join([p for p in parts if len(p) > 5])
        if len(candidate) > 20:
             print(f" Brute-force recovery success: {len(candidate)} chars")
             return _clean_text(candidate)
    except:
        pass

    return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    if DOCX_AVAILABLE:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            return _clean_text("\n".join(parts))
        except Exception as e:
            print(f"docx error: {e}")
    return ""


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a TXT/plain text file with encoding fallbacks."""
    for encoding in ('utf-8', 'utf-16', 'latin-1', 'cp1252'):
        try:
            return _clean_text(file_bytes.decode(encoding, errors='ignore'))
        except Exception:
            continue
    return ""


def extract_sections(text: str) -> dict:
    """
    Splits resume text into structured sections: Summary, Experience, Education, Skills, etc.
    Uses regex for improved section boundary detection.
    """
    sections = {
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
        "certifications": "",
    }
    
    if not text:
        return sections

    # Common section header patterns
    header_patterns = {
        "summary": r"(summary|objective|profile|about me|professional summary|overview)",
        "experience": r"(experience|work history|employment|professional background|career history)",
        "education": r"(education|academic|university|schooling|qualifications)",
        "skills": r"(skills|technical skills|expertise|competencies|technologies|tools)",
        "projects": r"(projects|personal projects|notable projects|portfolio)",
        "certifications": r"(certifications|certificates|licenses|credentials|training)",
    }

    lines = text.split('\n')
    current_section = "summary"
    buffers = {k: [] for k in sections.keys()}

    for line in lines:
        clean = line.strip()
        lower = clean.lower()
        if not clean: continue

        # Check if line is a header (short, and matches a pattern)
        is_header = False
        if len(clean) < 40:
            for sec, pat in header_patterns.items():
                # Anchored match or strong keyword match for short lines
                if re.search(rf'^({pat})[:\s]*$', lower) or (len(clean) < 25 and re.search(rf'\b{pat}\b', lower)):
                    current_section = sec
                    is_header = True
                    break
        
        if not is_header:
            buffers[current_section].append(clean)

    for sec in sections.keys():
        sections[sec] = "\n".join(buffers[sec]).strip()
    
    return sections


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Main entry point: dispatch by extension with smart fallbacks.
    Returns: cleaned plain text.
    """
    if not filename:
        filename = "unknown.txt"

    name_lower = filename.lower()
    base, ext = os.path.splitext(name_lower)
    text = ""

    if ext == '.pdf':
        text = extract_text_from_pdf(file_bytes)
        if not text or len(text.strip()) < 30:
            text = extract_text_from_docx(file_bytes) or text
    elif ext in ('.docx', '.doc'):
        text = extract_text_from_docx(file_bytes)
        if not text or len(text.strip()) < 30:
            text = extract_text_from_pdf(file_bytes) or text
    elif ext == '.txt':
        text = extract_text_from_txt(file_bytes)
    else:
        # Unknown extension  try fallbacks
        text = (
            extract_text_from_pdf(file_bytes) or
            extract_text_from_docx(file_bytes) or
            extract_text_from_txt(file_bytes)
        )

    if not text or len(text.strip()) < 10:
        print(f" Extraction failed for: {filename}")
        return ""

    print(f" Extracted {len(text)} chars from {filename}")
    return text
